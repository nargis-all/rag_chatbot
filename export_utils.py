from docx import Document
from io import BytesIO


def create_answer_docx(question, answer, context=None):
    """
    Generate a Word document containing the question, answer,
    and optionally the retrieved context, returned as bytes
    for direct download in Streamlit.
    """

    doc = Document()

    doc.add_heading("Customs RAG Chatbot - Answer", level=1)

    doc.add_heading("Question", level=2)
    doc.add_paragraph(question)

    doc.add_heading("Answer", level=2)
    doc.add_paragraph(answer)

    if context:
        doc.add_heading("Retrieved Context", level=2)
        doc.add_paragraph(context)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer